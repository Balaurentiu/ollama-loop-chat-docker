"""
Chat Export Module - Generate PDF/DOCX from chat messages with markdown support.
Supports dark and light themes for both output formats.
"""
import re
from io import BytesIO
from datetime import datetime


def generate_pdf(messages, theme='dark'):
    """
    Generate PDF from chat messages with markdown formatting.

    Args:
        messages: list of dicts with keys: role, text, timestamp
        theme: 'dark' or 'light'

    Returns:
        BytesIO buffer containing the PDF
    """
    import markdown
    from xhtml2pdf import pisa

    md = markdown.Markdown(extensions=['tables', 'fenced_code'])

    css = _get_pdf_css(theme)
    html_parts = [f'<html><head><meta charset="utf-8"><style>{css}</style></head><body>']
    html_parts.append('<h1 class="export-title">Chat Export</h1>')
    html_parts.append(f'<p class="export-meta">Exported: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")} | {len(messages)} message(s)</p>')
    html_parts.append('<hr>')

    for msg in messages:
        role_class = 'user' if msg['role'] == 'user' else 'agent'
        label = 'You' if msg['role'] == 'user' else 'Agent'
        md.reset()
        content_html = md.convert(msg.get('text', ''))
        timestamp = msg.get('timestamp', '')

        html_parts.append(f'''
        <div class="message {role_class}">
            <div class="msg-header">
                <span class="msg-role">{label}</span>
                <span class="msg-time">{timestamp}</span>
            </div>
            <div class="msg-content">{content_html}</div>
        </div>
        ''')

    html_parts.append('</body></html>')
    full_html = '\n'.join(html_parts)

    buffer = BytesIO()
    pisa.CreatePDF(full_html, dest=buffer)
    buffer.seek(0)
    return buffer


def generate_docx(messages, theme='dark'):
    """
    Generate DOCX from chat messages with markdown formatting.

    Args:
        messages: list of dicts with keys: role, text, timestamp
        theme: 'dark' or 'light'

    Returns:
        BytesIO buffer containing the DOCX
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    colors = _get_docx_colors(theme)

    if theme == 'dark':
        _set_docx_page_background(doc, colors['page_bg'])

    # Title
    title_para = doc.add_paragraph()
    run = title_para.add_run('Chat Export')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = colors['title']

    # Metadata
    meta_para = doc.add_paragraph()
    run = meta_para.add_run(f'Exported: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")} | {len(messages)} message(s)')
    run.font.size = Pt(10)
    run.font.color.rgb = colors['meta']

    doc.add_paragraph()  # spacer

    for msg in messages:
        label = 'You' if msg['role'] == 'user' else 'Agent'
        role_color = colors['user_label'] if msg['role'] == 'user' else colors['agent_label']
        timestamp = msg.get('timestamp', '')

        # Role + timestamp header
        header_para = doc.add_paragraph()
        role_run = header_para.add_run(label)
        role_run.bold = True
        role_run.font.size = Pt(11)
        role_run.font.color.rgb = role_color

        if timestamp:
            ts_run = header_para.add_run(f'  {timestamp}')
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = colors['timestamp']

        # Message content with markdown formatting
        _add_markdown_to_docx(doc, msg.get('text', ''), colors)

        doc.add_paragraph()  # spacer between messages

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def _get_pdf_font_face():
    """Return @font-face CSS for Unicode-capable DejaVu Sans font."""
    return '''
            @font-face {
                font-family: 'DejaVu';
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf');
            }
            @font-face {
                font-family: 'DejaVu';
                font-weight: bold;
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf');
            }
            @font-face {
                font-family: 'DejaVuMono';
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf');
            }
            @font-face {
                font-family: 'DejaVuMono';
                font-weight: bold;
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf');
            }
    '''


def _get_pdf_css(theme):
    """Return CSS string for PDF based on theme."""
    font_face = _get_pdf_font_face()
    if theme == 'dark':
        return font_face + '''
            @page { size: A4; margin: 1.5cm; background-color: #121212; }
            body { background-color: #121212; color: #e0e0e0; font-family: 'DejaVu', sans-serif; font-size: 11px; }
            .export-title { color: #7aa5d2; font-size: 20px; margin-bottom: 2px; }
            .export-meta { color: #888; font-size: 10px; margin-top: 0; }
            hr { border: none; border-top: 1px solid #444; }
            .message { margin-bottom: 14px; padding: 10px; border-radius: 6px; }
            .message.user { background-color: #0e639c; }
            .message.agent { background-color: #333333; }
            .msg-header { margin-bottom: 5px; }
            .msg-role { font-weight: bold; font-size: 11px; }
            .msg-time { font-size: 9px; color: #aaa; margin-left: 10px; }
            .msg-content { line-height: 1.5; }
            .msg-content code { background-color: rgba(0,0,0,0.3); padding: 1px 4px; font-family: 'DejaVuMono', monospace; font-size: 10px; }
            .msg-content pre { background-color: rgba(0,0,0,0.35); padding: 8px; font-family: 'DejaVuMono', monospace; font-size: 10px; white-space: pre-wrap; }
            .msg-content table { border-collapse: collapse; width: 100%; margin: 5px 0; }
            .msg-content th, .msg-content td { border: 1px solid #555; padding: 4px 8px; font-size: 10px; }
            .msg-content th { background-color: rgba(255,255,255,0.08); }
            .msg-content h1, .msg-content h2, .msg-content h3 { color: #7aa5d2; }
            .msg-content strong { color: #ffffff; }
            .msg-content blockquote { border-left: 3px solid #555; padding-left: 10px; color: #aaa; }
            .msg-content ul, .msg-content ol { padding-left: 20px; }
        '''
    else:
        return font_face + '''
            @page { size: A4; margin: 1.5cm; }
            body { background-color: #ffffff; color: #222222; font-family: 'DejaVu', sans-serif; font-size: 11px; }
            .export-title { color: #2c3e50; font-size: 20px; margin-bottom: 2px; }
            .export-meta { color: #888; font-size: 10px; margin-top: 0; }
            hr { border: none; border-top: 1px solid #ddd; }
            .message { margin-bottom: 14px; padding: 10px; border-radius: 6px; border: 1px solid #ddd; }
            .message.user { background-color: #e8f4fd; }
            .message.agent { background-color: #f5f5f5; }
            .msg-header { margin-bottom: 5px; }
            .msg-role { font-weight: bold; font-size: 11px; color: #333; }
            .msg-time { font-size: 9px; color: #999; margin-left: 10px; }
            .msg-content { line-height: 1.5; }
            .msg-content code { background-color: #f0f0f0; padding: 1px 4px; font-family: 'DejaVuMono', monospace; font-size: 10px; }
            .msg-content pre { background-color: #f5f5f5; padding: 8px; font-family: 'DejaVuMono', monospace; font-size: 10px; border: 1px solid #ddd; white-space: pre-wrap; }
            .msg-content table { border-collapse: collapse; width: 100%; margin: 5px 0; }
            .msg-content th, .msg-content td { border: 1px solid #ccc; padding: 4px 8px; font-size: 10px; }
            .msg-content th { background-color: #f0f0f0; }
            .msg-content h1, .msg-content h2, .msg-content h3 { color: #2c3e50; }
            .msg-content blockquote { border-left: 3px solid #ddd; padding-left: 10px; color: #666; }
            .msg-content ul, .msg-content ol { padding-left: 20px; }
        '''


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _get_docx_colors(theme):
    """Return color dict for DOCX based on theme."""
    from docx.shared import RGBColor
    if theme == 'dark':
        return {
            'page_bg': '121212',
            'title': RGBColor(0x7a, 0xa5, 0xd2),
            'meta': RGBColor(0x88, 0x88, 0x88),
            'text': RGBColor(0xe0, 0xe0, 0xe0),
            'user_label': RGBColor(0x4C, 0xAF, 0x50),
            'agent_label': RGBColor(0x21, 0x96, 0xF3),
            'timestamp': RGBColor(0xaa, 0xaa, 0xaa),
            'heading': RGBColor(0x7a, 0xa5, 0xd2),
            'bold': RGBColor(0xff, 0xff, 0xff),
            'code': RGBColor(0xce, 0x91, 0x78),
            'link': RGBColor(0x7a, 0xa5, 0xd2),
        }
    else:
        return {
            'page_bg': 'ffffff',
            'title': RGBColor(0x2c, 0x3e, 0x50),
            'meta': RGBColor(0x88, 0x88, 0x88),
            'text': RGBColor(0x22, 0x22, 0x22),
            'user_label': RGBColor(0x2e, 0x7d, 0x32),
            'agent_label': RGBColor(0x15, 0x65, 0xc0),
            'timestamp': RGBColor(0x99, 0x99, 0x99),
            'heading': RGBColor(0x2c, 0x3e, 0x50),
            'bold': RGBColor(0x00, 0x00, 0x00),
            'code': RGBColor(0xc7, 0x25, 0x4e),
            'link': RGBColor(0x15, 0x65, 0xc0),
        }


def _set_docx_page_background(doc, hex_color):
    """Set page background color for DOCX document."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    bg_elem = OxmlElement('w:background')
    bg_elem.set(qn('w:color'), hex_color)
    bg_elem.set(qn('w:themeColor'), 'dark1')
    doc.element.insert(0, bg_elem)


def _add_markdown_to_docx(doc, text, colors):
    """
    Parse markdown text and add formatted paragraphs to DOCX.
    Handles: headings, bold, italic, code inline, code blocks, lists, links.
    """
    from docx.shared import Pt

    lines = text.split('\n')
    in_code_block = False
    code_block_lines = []

    for line in lines:
        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_para = doc.add_paragraph()
                code_text = '\n'.join(code_block_lines)
                code_run = code_para.add_run(code_text)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = colors['code']
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # Headings
        header_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if header_match:
            level = len(header_match.group(1))
            para = doc.add_paragraph()
            run = para.add_run(header_match.group(2))
            run.bold = True
            run.font.size = Pt(max(10, 18 - (level * 2)))
            run.font.color.rgb = colors['heading']
            continue

        # List items (unordered and ordered)
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)', line)
        if list_match:
            indent = len(list_match.group(1))
            content = list_match.group(3)
            bullet = list_match.group(2)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(20 + indent * 10)
            # Add bullet/number prefix
            prefix = '\u2022 ' if bullet in ('-', '*', '+') else f'{bullet} '
            prefix_run = para.add_run(prefix)
            prefix_run.font.size = Pt(10)
            prefix_run.font.color.rgb = colors['text']
            _add_inline_formatting(para, content, colors)
            continue

        # Empty lines - skip
        if not line.strip():
            continue

        # Regular paragraph with inline formatting
        para = doc.add_paragraph()
        _add_inline_formatting(para, line, colors)


def _add_inline_formatting(paragraph, text, colors):
    """
    Add inline markdown formatting to a paragraph.
    Handles: **bold**, *italic*, `code`, [text](url)
    """
    from docx.shared import Pt

    # Split text by inline markdown patterns
    # Order matters: **bold** before *italic*
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = colors.get('bold', colors['text'])
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.color.rgb = colors['text']
            run.font.size = Pt(10)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = colors['code']
        elif part.startswith('['):
            link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if link_match:
                run = paragraph.add_run(link_match.group(1))
                run.font.color.rgb = colors['link']
                run.underline = True
                run.font.size = Pt(10)
            else:
                run = paragraph.add_run(part)
                run.font.color.rgb = colors['text']
                run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.color.rgb = colors['text']
            run.font.size = Pt(10)
